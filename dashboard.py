import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import os

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx import Document

# =========================
# PAGE CONFIG
# =========================
st.set_page_config("Universal Data Analyzer",  layout="wide")

st.markdown("<h1 style='text-align:center;'>Universal Data Analyzer</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center;'>Business & Data Analyst Project</p>", unsafe_allow_html=True)

# =========================
# FILE UPLOAD (ALL FILES)
# =========================
uploaded_file = st.file_uploader(
    " Upload Dataset (CSV, Excel, TSV, TXT, JSON)",
    type=["csv", "xlsx", "xls", "tsv", "txt", "json"]
)

# =========================
# FILE LOADER
# =========================
def load_file(file):
    ext = file.name.split(".")[-1].lower()
    if ext == "csv":
        return pd.read_csv(file)
    elif ext in ["xlsx", "xls"]:
        return pd.read_excel(file)
    elif ext == "tsv":
        return pd.read_csv(file, sep="\t")
    elif ext == "txt":
        return pd.read_csv(file, sep=None, engine="python")
    elif ext == "json":
        return pd.read_json(file)
    else:
        return None

if uploaded_file:

    df = load_file(uploaded_file)

    if df is None:
        st.stop()

    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    categorical_cols = df.select_dtypes(exclude=np.number).columns.tolist()

    # =========================
    # DATA PREVIEW
    # =========================
    st.subheader(" Dataset Preview")
    st.dataframe(df.head(), use_container_width=True)

    # =========================
    # FILTERS
    # =========================
    st.sidebar.header("Filters")
    if categorical_cols:
        cat_filter = st.sidebar.selectbox("Filter Column", categorical_cols)
        selected_vals = st.sidebar.multiselect(
            "Select Values", df[cat_filter].unique()
        )
        if selected_vals:
            df = df[df[cat_filter].isin(selected_vals)]

    # =========================
    # KPIs
    # =========================
    st.subheader(" Key Metrics")

    num_col = st.selectbox("Select Numeric Column", numeric_cols)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("SUM", round(df[num_col].sum(), 2))
    c2.metric("MEAN", round(df[num_col].mean(), 2))
    c3.metric("MIN", round(df[num_col].min(), 2))
    c4.metric("MAX", round(df[num_col].max(), 2))

    # =========================
    # LINE CHART
    # =========================
    line_fig = px.line(df, y=num_col, markers=True, title="Line Chart")
    line_fig.write_image("line.png", engine="kaleido")
    st.plotly_chart(line_fig, use_container_width=True)

    # =========================
    # BAR / PIE / DONUT
    # =========================
    cat_col = st.selectbox("Select Categorical Column", categorical_cols)

    vc = df[cat_col].value_counts().reset_index()
    vc.columns = [cat_col, "Count"]

    bar_fig = px.bar(vc, x=cat_col, y="Count", title="Bar Chart")
    pie_fig = px.pie(vc, names=cat_col, values="Count", title="Pie Chart")
    donut_fig = px.pie(vc, names=cat_col, values="Count", hole=0.5, title="Donut Chart")

    bar_fig.write_image("bar.png", engine="kaleido")
    pie_fig.write_image("pie.png", engine="kaleido")
    donut_fig.write_image("donut.png", engine="kaleido")

    st.plotly_chart(bar_fig, use_container_width=True)
    st.plotly_chart(pie_fig, use_container_width=True)
    st.plotly_chart(donut_fig, use_container_width=True)

    # =========================
    # CLUSTERED CHARTS
    # =========================
    grouped = df.groupby(cat_col)[num_col].mean().reset_index()

    cluster_col = px.bar(grouped, x=cat_col, y=num_col, title="Clustered Column Chart")
    cluster_bar = px.bar(grouped, y=cat_col, x=num_col, orientation="h", title="Clustered Bar Chart")

    cluster_col.write_image("cluster_col.png", engine="kaleido")
    cluster_bar.write_image("cluster_bar.png", engine="kaleido")

    st.plotly_chart(cluster_col, use_container_width=True)
    st.plotly_chart(cluster_bar, use_container_width=True)

    # =========================
    # CORRELATION
    # =========================
    if len(numeric_cols) > 1:
        corr = df[numeric_cols].corr()
        heat = px.imshow(corr, text_auto=True, title="Correlation Heatmap")
        heat.write_image("corr.png", engine="kaleido")
        st.plotly_chart(heat, use_container_width=True)

    # =========================
    # INSIGHTS & SUMMARY
    # =========================
    st.subheader("Insights")

    insights = [
        f"Average {num_col} is {round(df[num_col].mean(),2)}",
        f"Maximum value is {df[num_col].max()}",
        f"Most frequent category in {cat_col} is '{df[cat_col].mode()[0]}'"
    ]

    for i in insights:
        st.write( i)

    summary = f"""
This dataset contains {df.shape[0]} records and {df.shape[1]} columns.
The analysis includes KPIs, trend analysis, category distribution,
clustered comparisons, and correlation analysis to support
business decision-making.
"""

    st.subheader("Summary")
    st.write(summary)

    # =========================
    # PDF REPORT
    # =========================
    def create_pdf():
        doc = SimpleDocTemplate("Business_Report.pdf", pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph("Universal Data Analyzer – Business Report", styles["Title"]), Spacer(1, 12)]

        for i in insights:
            story.append(Paragraph(i, styles["Normal"]))

        for img in ["line.png","bar.png","pie.png","donut.png","cluster_col.png","cluster_bar.png","corr.png"]:
            if os.path.exists(img):
                story.append(Image(img, 5*inch, 3*inch))
                story.append(Spacer(1, 12))

        story.append(Paragraph(summary, styles["Normal"]))
        doc.build(story)
        return "Business_Report.pdf"

    # =========================
    # WORD REPORT
    # =========================
    def create_word():
        doc = Document()
        doc.add_heading("Universal Data Analyzer – Analyst Report", 1)
        for i in insights:
            doc.add_paragraph(i)
        doc.add_paragraph(summary)

        for img in ["line.png","bar.png","pie.png","donut.png","cluster_col.png","cluster_bar.png","corr.png"]:
            if os.path.exists(img):
                doc.add_picture(img, width=inch*5)

        doc.save("Analyst_Report.docx")
        return "Analyst_Report.docx"

    st.subheader("Download Reports")

    if st.button(" Download PDF"):
        with open(create_pdf(), "rb") as f:
            st.download_button("Download PDF", f)

    if st.button(" Download Word"):
        with open(create_word(), "rb") as f:
            st.download_button("Download Word", f)

else:
    st.info("Upload any dataset file to begin analysis")
